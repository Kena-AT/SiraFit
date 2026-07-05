"""Resume export services: HTML and DOCX download generation."""

import json
import io
from typing import Optional

from app.models.job import ResumeVersion
from app.services.resume_generation import render_resume_html


def export_resume_html(version: ResumeVersion) -> str:
    """Render a ResumeVersion's content as a standalone HTML string."""
    data = json.loads(version.content) if version.content else {}
    template = version.template or "minimal"
    return render_resume_html(data, template)


def export_resume_docx(version: ResumeVersion) -> io.BytesIO:
    """Generate a .docx file from a ResumeVersion's content.

    Returns a BytesIO stream suitable for use as a file download.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = json.loads(version.content) if version.content else {}
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Name
    name = data.get("name", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)

    # Contact
    contact_parts = [
        data.get("email", ""),
        data.get("phone", ""),
        data.get("location", ""),
        data.get("linkedin", ""),
        data.get("github", ""),
        data.get("website", ""),
    ]
    contact = "  |  ".join([c for c in contact_parts if c])
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact)
        run.font.size = Pt(9)
        run.font.color.rgb = None  # default

    # Summary
    if data.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(data["summary"])

    # Experience
    if data.get("experience"):
        doc.add_heading("Experience", level=2)
        for exp in data["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            run.bold = True
            run.font.size = Pt(11)
            if exp.get("period"):
                run = p.add_run(f"   {exp['period']}")
                run.font.size = Pt(10)
            if exp.get("location"):
                p2 = doc.add_paragraph(exp["location"])
                p2.style = doc.styles["Normal"]
                p2.paragraph_format.space_before = Pt(0)
            if exp.get("bullets"):
                for bullet in exp["bullets"]:
                    doc.add_paragraph(bullet, style="List Bullet")

    # Projects
    if data.get("projects"):
        doc.add_heading("Projects", level=2)
        for proj in data["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", ""))
            run.bold = True
            run.font.size = Pt(11)
            if proj.get("description"):
                doc.add_paragraph(proj["description"])

    # Skills
    if data.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(" · ".join(data["skills"]))

    # Education
    if data.get("education"):
        doc.add_heading("Education", level=2)
        for edu in data["education"]:
            parts = [edu.get("institution", "")]
            if edu.get("degree"):
                parts.append(edu["degree"])
            if edu.get("field_of_study"):
                parts.append(edu["field_of_study"])
            if edu.get("period"):
                parts.append(edu["period"])
            doc.add_paragraph(" — ".join(parts))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
